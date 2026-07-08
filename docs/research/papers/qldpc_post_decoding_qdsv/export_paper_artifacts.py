"""Export the qLDPC/QDSV paper draft to DOCX and LaTeX.

The source of truth is ``paper_draft.md``. This exporter intentionally keeps
the output conservative: formal Word styling, real Word tables, and a
self-contained LaTeX article file suitable for manual polishing.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "paper_draft.md"
EXPORT_DIR = ROOT / "exports"
DOCX_OUT = EXPORT_DIR / "QDSV_QIntent_qLDPC_decision_governance_paper.docx"
TEX_OUT = EXPORT_DIR / "QDSV_QIntent_qLDPC_decision_governance_paper.tex"

DOWNLOADS = Path.home() / "Downloads"
DOCX_DOWNLOAD = DOWNLOADS / DOCX_OUT.name
TEX_DOWNLOAD = DOWNLOADS / TEX_OUT.name


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = styles.add_style("QDSV Code", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)


def _add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_width(table)
    _set_cell_margins(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx].strip() if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _add_inline_runs(p, text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx == 0:
                _set_cell_shading(cell, "F2F4F7")
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def _parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build_docx(markdown: str) -> None:
    doc = Document()
    _style_document(doc)
    lines = markdown.splitlines()
    in_code = False
    code_lines: list[str] = []
    i = 0
    title_added = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                for code_line in code_lines:
                    doc.add_paragraph(code_line, style="QDSV Code")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = _parse_markdown_table(lines, i)
            _add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(title)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(11, 37, 69)
            title_added = True
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)
        i += 1

    if not title_added:
        doc.add_heading("QDSV/QIntent qLDPC Decision Governance Paper", level=0)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "QDSV/QIntent qLDPC decision-governance draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(DOCX_OUT)


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    out = "".join(replacements.get(ch, ch) for ch in text)
    out = out.replace("->", r"$\rightarrow$")
    return out


def latex_inline(text: str) -> str:
    parts = re.split(r"(`[^`]+`)", text)
    rendered = []
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            rendered.append(r"\texttt{" + latex_escape(part[1:-1]) + "}")
        else:
            rendered.append(latex_escape(part))
    return "".join(rendered)


def _latex_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    cols = max(len(row) for row in rows)
    col_width = max(0.10, min(0.32, 0.92 / cols))
    col_spec = f">{{\\raggedright\\arraybackslash}}p{{{col_width:.3f}\\textwidth}}" * cols
    out = ["\\begin{small}", f"\\begin{{longtable}}{{{col_spec}}}", "\\toprule"]
    for idx, row in enumerate(rows):
        padded = row + [""] * (cols - len(row))
        out.append(" & ".join(latex_inline(cell.strip()) for cell in padded) + r" \\")
        if idx == 0:
            out.append("\\midrule")
    out.extend(["\\bottomrule", "\\end{longtable}", "\\end{small}", ""])
    return "\n".join(out)


def build_latex(markdown: str) -> None:
    lines = markdown.splitlines()
    title = "QDSV/QIntent as a Guarded Semantic Decision Layer for LDPC/qLDPC-Style Post-Decoding Correction Governance"
    body: list[str] = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{hyperref}",
        r"\usepackage{booktabs}",
        r"\usepackage{longtable}",
        r"\usepackage{array}",
        r"\usepackage{xcolor}",
        r"\usepackage{enumitem}",
        r"\usepackage{listings}",
        r"\lstset{basicstyle=\ttfamily\small,breaklines=true,frame=single,columns=fullflexible}",
        r"\hypersetup{colorlinks=true,linkcolor=blue,urlcolor=blue,citecolor=blue}",
        "",
        r"\title{" + latex_escape(title) + "}",
        r"\author{QDSV/QIntent Research Draft}",
        r"\date{\today}",
        "",
        r"\begin{document}",
        r"\maketitle",
        "",
    ]
    in_code = False
    code_lines: list[str] = []
    in_itemize = False
    in_enumerate = False
    i = 0

    def close_lists() -> None:
        nonlocal in_itemize, in_enumerate
        if in_itemize:
            body.append(r"\end{itemize}")
            in_itemize = False
        if in_enumerate:
            body.append(r"\end{enumerate}")
            in_enumerate = False

    def clean_heading(text: str) -> str:
        return re.sub(r"^\d+(?:\.\d+)?\.?\s+", "", text.strip())

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            close_lists()
            if in_code:
                body.append(r"\begin{lstlisting}")
                body.extend(code_lines)
                body.append(r"\end{lstlisting}")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            close_lists()
            body.append("")
            i += 1
            continue
        if stripped.startswith("|"):
            close_lists()
            rows, i = _parse_markdown_table(lines, i)
            body.append(_latex_table(rows))
            continue
        if stripped.startswith("# "):
            # Title already handled.
            i += 1
            continue
        if stripped.startswith("## "):
            close_lists()
            body.append(r"\section{" + latex_inline(clean_heading(stripped[3:].strip())) + "}")
        elif stripped.startswith("### "):
            close_lists()
            body.append(r"\subsection{" + latex_inline(clean_heading(stripped[4:].strip())) + "}")
        elif re.match(r"^\d+\.\s+", stripped):
            if in_itemize:
                body.append(r"\end{itemize}")
                in_itemize = False
            if not in_enumerate:
                body.append(r"\begin{enumerate}[leftmargin=*]")
                in_enumerate = True
            body.append(r"\item " + latex_inline(re.sub(r"^\d+\.\s+", "", stripped)))
        elif stripped.startswith("- "):
            if in_enumerate:
                body.append(r"\end{enumerate}")
                in_enumerate = False
            if not in_itemize:
                body.append(r"\begin{itemize}[leftmargin=*]")
                in_itemize = True
            body.append(r"\item " + latex_inline(stripped[2:].strip()))
        else:
            close_lists()
            body.append(latex_inline(stripped) + "\n")
        i += 1
    close_lists()
    body.extend(["", r"\end{document}", ""])
    TEX_OUT.write_text("\n".join(body), encoding="utf-8")


def main() -> None:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = SOURCE.read_text(encoding="utf-8")
    build_docx(markdown)
    build_latex(markdown)
    shutil.copy2(DOCX_OUT, DOCX_DOWNLOAD)
    shutil.copy2(TEX_OUT, TEX_DOWNLOAD)
    print(f"DOCX: {DOCX_OUT}")
    print(f"TEX:  {TEX_OUT}")
    print(f"DOCX copy: {DOCX_DOWNLOAD}")
    print(f"TEX copy:  {TEX_DOWNLOAD}")


if __name__ == "__main__":
    main()
